from io import BytesIO
from django.http import HttpResponse
from django.utils import timezone
from docx import Document
from .models import Contract


def export_contract_analysis_docx(contract: Contract) -> HttpResponse:
    doc = Document()
    doc.add_heading(f"Vertragsanalyse: {contract.title}", level=1)
    doc.add_paragraph(f"Erstellt: {timezone.now().strftime('%Y-%m-%d %H:%M')}")
    if contract.ai_summary:
        doc.add_heading("Zusammenfassung", level=2)
        doc.add_paragraph(contract.ai_summary)
    if contract.ai_risks:
        doc.add_heading("Risiken", level=2)
        for item in contract.ai_risks:
            doc.add_paragraph(str(item), style='List Bullet')
    if contract.ai_checklist:
        doc.add_heading("Checkliste", level=2)
        for item in contract.ai_checklist:
            doc.add_paragraph(str(item), style='List Bullet')
    if contract.ai_key_dates:
        doc.add_heading("Wichtige Termine", level=2)
        for item in contract.ai_key_dates:
            label = item.get('label', '') if isinstance(item, dict) else ''
            date = item.get('date', '') if isinstance(item, dict) else ''
            note = item.get('note', '') if isinstance(item, dict) else ''
            line = f"{label}"
            if date:
                line += f" ({date})"
            if note:
                line += f" - {note}"
            doc.add_paragraph(line, style='List Bullet')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"contract_analysis_{contract.id}.docx"
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response
